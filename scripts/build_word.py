# -*- coding: utf-8 -*-
"""把一个课程笔记目录里的 Markdown 合并成"便于纸质快速查找"的 Word 文档。

为什么要这个脚本：开卷考查的是"翻得快"。纸质版没有 Ctrl+F，所以成品必须
做这些事——前置索引按拼音排序（像字典）、每讲另起一页、页眉显示当前讲名、
表头跨页重复、页脚页码。这些在 Word 里手动调很费劲，脚本一次做好。

用法：
    python build_word.py <课程笔记目录>            # 自动扫描
    python build_word.py <目录> --out 成品.docx
若目录下有 _manifest.json，则按它指定的标题/顺序/排序来编译（推荐，
由 skill 的"建规则"阶段生成）；没有则回退到：00_* 文件在前并按拼音排序
表格、其余文件按文件名顺序。

manifest 里 "mode":"closed" 切换成闭卷背诵手册；文件 role 决定渲染：
flashcard=答案列加底纹+对折提示、selftest=另起页加"先做后看"提示、
memorize/unit=正常渲染。开卷（默认/mode 缺省）行为不变。

依赖：python-docx；若有需排序的索引表还需 pypinyin。
Windows 提示：若目标 docx 正在 Word 中打开会锁文件，脚本会自动改存
"<名>_新.docx"。
"""
import io
import json
import os
import re
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

CN_BODY = "宋体"
CN_HEAD = "黑体"

# ---------- LaTeX 数学 → OMML（Word 原生公式渲染） ----------
# 笔记里大量 $...$ / $$...$$ 公式。纯文本直接落进 Word 会显示成字面 "$x^2$"。
# 这里把 LaTeX 先用 latex2mathml 转 MathML，再用 Office 自带的 MML2OMML.XSL
# 转成 OMML（Word 原生公式标记），插进段落即可被 Word 正常渲染。
# 缺依赖（latex2mathml / 找不到 XSL）时自动回退为原文，保证脚本仍能产出。
_OMML_XFORM = None
_OMML_READY = None


def _get_omml_transform():
    global _OMML_XFORM, _OMML_READY
    if _OMML_READY is not None:
        return _OMML_XFORM
    try:
        import glob
        from lxml import etree
        cand = (glob.glob(r'C:\Program Files*\Microsoft Office\root\Office*\MML2OMML.XSL')
                + glob.glob(r'C:\Program Files*\Microsoft Office\Office*\MML2OMML.XSL'))
        import latex2mathml.converter  # noqa: F401  仅探测可用性
        if not cand:
            _OMML_READY = False
            sys.stderr.write('[警告] 未找到 MML2OMML.XSL，公式将以原文回退显示。\n')
            return None
        _OMML_XFORM = etree.XSLT(etree.parse(cand[0]))
        _OMML_READY = True
    except Exception as e:
        _OMML_READY = False
        _OMML_XFORM = None
        sys.stderr.write('[警告] LaTeX→OMML 不可用（%s），公式以原文回退。'
                         '装一下：pip install latex2mathml\n' % e)
    return _OMML_XFORM


def latex_to_omml(latex):
    """LaTeX 片段 → docx 可插入的 <m:oMath> 元素；失败返回 None。"""
    try:
        import latex2mathml.converter
        from lxml import etree
        xform = _get_omml_transform()
        if xform is None:
            return None
        # 规避 Office MML2OMML.XSL 的方向定界符 bug：\lVert/\rVert/\lvert/\rvert
        # （带方向）会被 XSLT 当成对的伸缩定界符 <m:d>，紧跟下标/上标时内容被吞成空
        # <m:e/>（如 \lVert x\rVert_2 只剩 "2"）。换成无方向的 \Vert/\vert，渲染成普通
        # ‖ / | 算符、内容完整保留，视觉一致。
        latex = (latex.replace(r'\lVert', r'\Vert').replace(r'\rVert', r'\Vert')
                      .replace(r'\lvert', r'\vert').replace(r'\rvert', r'\vert'))
        mml = latex2mathml.converter.convert(latex)
        omml = xform(etree.fromstring(mml))
        return parse_xml(etree.tostring(omml).decode('utf-8'))
    except Exception:
        return None

# ---------- 底层工具 ----------

def set_cn_font(run, name=CN_BODY, size=None, bold=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for attr in ('w:eastAsia', 'w:ascii', 'w:hAnsi'):
        rFonts.set(qn(attr), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color is not None: run.font.color.rgb = color


def shade(elem_pr, hexcolor):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    elem_pr.append(shd)


def shade_cell(cell, hexcolor):
    shade(cell._tc.get_or_add_tcPr(), hexcolor)


def shade_para(p, hexcolor):
    shade(p._p.get_or_add_pPr(), hexcolor)


def set_table_borders(table):
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), '999999')
        borders.append(e)
    table._tbl.tblPr.append(borders)


def repeat_header_row(row):
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true')
    row._tr.get_or_add_trPr().append(th)


def add_field(paragraph, instr):
    run = paragraph.add_run()
    for typ in ('begin',):
        f = OxmlElement('w:fldChar'); f.set(qn('w:fldCharType'), typ); run._r.append(f)
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr; run._r.append(it)
    for typ in ('separate', 'end'):
        f = OxmlElement('w:fldChar'); f.set(qn('w:fldCharType'), typ); run._r.append(f)
    return run


def add_inline(paragraph, text, size=10.5, bold=False):
    """处理 $行内/行间公式$（转 Word 原生 OMML）、**加粗** 与 `代码`（去反引号）。"""
    text = (text or '').replace('`', '')
    # 兼容 LaTeX 定界符 \(...\) / \[...\]：统一成 $...$ / $$...$$，否则会原样落成字面文本
    text = re.sub(r'\\\[(.+?)\\\]', r'$$\1$$', text, flags=re.S)
    text = re.sub(r'\\\((.+?)\\\)', r'$\1$', text)
    # 先切出公式段：$$...$$ 优先于 $...$；公式内不跨行
    for seg in re.split(r'(\$\$.+?\$\$|\$[^$\n]+?\$)', text):
        if not seg:
            continue
        if len(seg) >= 2 and seg[0] == '$' and seg[-1] == '$':
            tex = seg[2:-2] if seg.startswith('$$') else seg[1:-1]
            omml = latex_to_omml(tex.strip())
            if omml is not None:
                paragraph._p.append(omml)   # 插入 Word 原生公式
                continue
            # 回退：转换失败保留原文，至少不丢内容
            set_cn_font(paragraph.add_run(tex), CN_BODY, size, bold)
            continue
        for part in re.split(r'(\*\*.*?\*\*)', seg):
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                set_cn_font(paragraph.add_run(part[2:-2]), CN_BODY, size, True)
            else:
                set_cn_font(paragraph.add_run(part), CN_BODY, size, bold)


# ---------- 拼音排序 ----------

def pinyin_key(cell):
    s = re.sub(r'\*\*', '', cell or '').strip()
    s = re.sub(r'^[《〈「『“”\"\'（()【\[\s·]+', '', s)
    try:
        from pypinyin import lazy_pinyin
        return ''.join(lazy_pinyin(s)).lower()
    except ImportError:
        return s.lower()


# ---------- Markdown 解析 ----------

def is_table_sep(line):
    return bool(re.match(r'^\s*\|[\s:|-]+\|\s*$', line)) and '-' in line


def split_row(line):
    line = line.strip()
    if line.startswith('|'): line = line[1:]
    if line.endswith('|'): line = line[:-1]
    return [c.strip() for c in line.split('|')]


def render_table(doc, header, rows, sort_rows=False, flashcard=False):
    """flashcard=True 时把最后一列（答案列）单元格加浅底纹，便于打印后对折/盖住自测。"""
    if sort_rows:
        rows = sorted(rows, key=lambda r: pinyin_key(r[0] if r else ''))
    ncol = len(header)
    ans_col = ncol - 1 if (flashcard and ncol >= 2) else None
    table = doc.add_table(rows=1, cols=ncol)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_table_borders(table)
    hdr = table.rows[0]
    for i, cell in enumerate(hdr.cells):
        shade_cell(cell, 'D9E2F3')
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(p, header[i] if i < len(header) else '', 9.5, True)
    repeat_header_row(hdr)
    for r in rows:
        cells = table.add_row().cells
        for i in range(ncol):
            if i == ans_col:
                shade_cell(cells[i], 'EDEDED')  # 答案列底纹，遮挡自测
            add_inline(cells[i].paragraphs[0], r[i] if i < len(r) else '', 9.5)
    doc.add_paragraph()


def render_markdown(doc, lines, sort_tables=False, flashcard=False):
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]; s = line.strip()
        if s.startswith('|') and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(line); rows = []; i += 2
            while i < n and lines[i].strip().startswith('|'):
                rows.append(split_row(lines[i])); i += 1
            if flashcard:  # 每张闪卡表前提示对折遮挡
                hint = doc.add_paragraph(); hint.paragraph_format.space_after = Pt(2)
                set_cn_font(hint.add_run('✂ 对折或用手盖住右列「答案」，先看左列自测'),
                            CN_BODY, 9, color=RGBColor(0x80, 0x60, 0x00))
            render_table(doc, header, rows, sort_tables, flashcard); continue
        for lvl, mark, size in ((3, '### ', 12), (2, '## ', 14), (1, '# ', 18)):
            if s.startswith(mark):
                p = doc.add_heading(level=lvl)
                for r in list(p.runs): r.text = ''
                if lvl == 1:
                    p.paragraph_format.page_break_before = True  # 每讲另起一页
                add_inline(p, s[len(mark):], size, True)
                i += 1; break
        else:
            if s in ('---', '***', '___'):
                i += 1; continue
            if s.startswith('>'):
                ql = []
                while i < n and lines[i].strip().startswith('>'):
                    ql.append(lines[i].strip()[1:].strip()); i += 1
                p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(6)
                shade_para(p, 'F2F2F2')
                for j, seg in enumerate('\n'.join(ql).split('\n')):
                    if j > 0: p.add_run().add_break()
                    add_inline(p, seg, 10)
                continue
            m = re.match(r'^(\s*)[-*]\s+(.*)$', line)
            if m:
                level = len(m.group(1)) // 2
                p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
                p.paragraph_format.left_indent = Cm(0.6 + 0.5 * level)
                p.paragraph_format.space_after = Pt(1)
                add_inline(p, m.group(2), 10.5); i += 1; continue
            if s == '':
                i += 1; continue
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
            add_inline(p, s, 10.5); i += 1


# ---------- 文档框架 ----------

def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = CN_BODY; normal.font.size = Pt(10.5)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), CN_BODY)
    for lvl, size in (('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12)):
        st = doc.styles[lvl]
        st.font.name = CN_HEAD; st.font.size = Pt(size)
        st.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        st.element.rPr.rFonts.set(qn('w:eastAsia'), CN_HEAD)


def setup_page(doc):
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for m in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(sec, m, Cm(1.8))
    sec.different_first_page_header_footer = True  # 封面不显示页眉页码
    hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # 中文版 Word 必须用本地化样式名"标题 1"，英文 "Heading 1" 会满屏报错
    add_field(hp, 'STYLEREF "标题 1" \\* MERGEFORMAT')
    for r in hp.runs: set_cn_font(r, CN_BODY, 9, color=RGBColor(0x80, 0x80, 0x80))
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cn_font(fp.add_run('— '), CN_BODY, 9, color=RGBColor(0x80, 0x80, 0x80))
    add_field(fp, 'PAGE')
    set_cn_font(fp.add_run(' —'), CN_BODY, 9, color=RGBColor(0x80, 0x80, 0x80))
    for r in fp.runs: set_cn_font(r, CN_BODY, 9, color=RGBColor(0x80, 0x80, 0x80))


def add_cover(doc, title, subtitle, blurb, steps):
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cn_font(p.add_run(title), CN_HEAD, 34, True, RGBColor(0x1F, 0x38, 0x64))
    if subtitle:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cn_font(p.add_run(subtitle), CN_HEAD, 20, True, RGBColor(0x1F, 0x38, 0x64))
    if blurb:
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cn_font(p.add_run(blurb), CN_BODY, 12, color=RGBColor(0x80, 0x80, 0x80))
    if steps:
        doc.add_paragraph(); doc.add_paragraph()
        p = doc.add_paragraph(); shade_para(p, 'FFF2CC'); p.paragraph_format.space_before = Pt(6)
        set_cn_font(p.add_run('📖 怎么用这本手册'), CN_HEAD, 14, True, RGBColor(0x80, 0x60, 0x00))
        for item in steps:
            t, b = (item.get('title', ''), item.get('body', '')) if isinstance(item, dict) else ('', item)
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
            if t: set_cn_font(p.add_run('▸ ' + t + '：'), CN_HEAD, 11, True, RGBColor(0x1F, 0x38, 0x64))
            set_cn_font(p.add_run(b), CN_BODY, 10.5)


# ---------- 主流程 ----------

def discover(folder):
    """无 manifest 时的回退：00_* 在前并排序表格，其余按名排序。"""
    # 排除内部文件：规则、manifest、MEMORY 不进成品手册
    skip = ('整理规则', '规则', 'rules', 'readme', 'memory')
    files = [f for f in os.listdir(folder) if f.lower().endswith('.md')
             and not f.startswith('_')
             and not any(k in f.lower() for k in skip)]
    index = sorted([f for f in files if f.startswith('00')])
    units = sorted([f for f in files if not f.startswith('00')])
    spec = [{'path': f, 'sort_tables': True} for f in index]
    spec += [{'path': f, 'sort_tables': False} for f in units]
    return spec


def main():
    args = [a for a in sys.argv[1:] if not a.startswith('--')]
    folder = args[0] if args else '.'
    out = None
    if '--out' in sys.argv:
        out = sys.argv[sys.argv.index('--out') + 1]

    manifest_path = os.path.join(folder, '_manifest.json')
    manifest = {}
    if os.path.exists(manifest_path):
        with open(manifest_path, encoding='utf-8') as f:
            manifest = json.load(f)

    closed = manifest.get('mode') == 'closed'
    default_sub = '期末背诵手册' if closed else '期末速查手册'

    doc = Document(); setup_styles(doc); setup_page(doc)
    add_cover(doc,
              manifest.get('title', os.path.basename(os.path.abspath(folder))),
              manifest.get('subtitle', default_sub),
              manifest.get('blurb', ''),
              manifest.get('cover_steps', []))

    file_spec = manifest.get('files') or discover(folder)
    for item in file_spec:
        path = os.path.join(folder, item['path'])
        if not os.path.exists(path):
            print('  [缺文件，跳过]', item['path']); continue
        role = item.get('role', '')
        if role == 'selftest':  # 自测题：提醒先做后看
            p = doc.add_paragraph(); p.paragraph_format.page_break_before = True
            p.paragraph_format.space_after = Pt(2)
            set_cn_font(p.add_run('⚠ 先独立做题、不要看答案，做完再翻到末尾「答案」核对。'),
                        CN_BODY, 10, True, RGBColor(0x80, 0x60, 0x00))
        with open(path, encoding='utf-8') as f:
            render_markdown(doc, f.read().split('\n'),
                            item.get('sort_tables', False),
                            flashcard=(role == 'flashcard'))

    if not out:
        out = os.path.join(folder, manifest.get('title', '速查手册') + '.docx')
    try:
        doc.save(out)
    except PermissionError:
        base, ext = os.path.splitext(out); out = base + '_新' + ext; doc.save(out)
    print('SAVED:', out)


if __name__ == '__main__':
    main()
